import os
import cv2
import pytesseract
import speech_recognition as sr
from pydub import AudioSegment
from pydub.utils import make_chunks
from moviepy.editor import VideoFileClip
from docx import Document

# Set Tesseract OCR path if required
# Uncomment and set the path to your Tesseract executable
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# 1. Audio to Text Conversion
def audio_to_text(video_file, chunk_size=10000):
    """
    Extracts audio from a video file and converts it to text.
    """
    print("Extracting audio...")
    audio = AudioSegment.from_file(video_file)
    chunks = make_chunks(audio, chunk_size)
    recognizer = sr.Recognizer()
    full_text = []

    for i, chunk in enumerate(chunks):
        chunk_file = f"chunk_{i}.wav"
        chunk.export(chunk_file, format="wav")
        with sr.AudioFile(chunk_file) as source:
            audio_data = recognizer.record(source)
            try:
                text = recognizer.recognize_google(audio_data)
                full_text.append(text)
            except sr.UnknownValueError:
                full_text.append("[Unintelligible segment]")
            except sr.RequestError as e:
                full_text.append(f"[Error: {e}]")
        os.remove(chunk_file)

    return " ".join(full_text)

# 2. Image to Text Conversion (OCR)
def video_to_text(video_file, frame_skip=30):
    """
    Extracts text from video frames using OCR.
    """
    print("Extracting text from video frames...")
    cap = cv2.VideoCapture(video_file)
    frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    frame_texts = []

    for i in range(frame_count):
        ret, frame = cap.read()
        if not ret or i % frame_skip != 0:
            continue
        # Convert frame to grayscale
        gray_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        # Perform OCR
        text = pytesseract.image_to_string(gray_frame)
        if text.strip():
            frame_texts.append(f"Frame {i}: {text.strip()}")
    cap.release()

    return "\n".join(frame_texts)

# 3. Document Compilation
def create_document(audio_text, frame_text, output_file="output.docx"):
    """
    Combines extracted audio and frame text into a single DOCX document.
    """
    print("Compiling document...")
    doc = Document()
    doc.add_heading("Extracted Text from Video", level=1)

    doc.add_heading("Audio to Text Conversion", level=2)
    doc.add_paragraph(audio_text)

    doc.add_heading("Image to Text Conversion (OCR)", level=2)
    doc.add_paragraph(frame_text)

    doc.save(output_file)
    print(f"Document saved as {output_file}")

# Main Function
if __name__ == "__main__":
    video_file = "your_video_file.mp4"  # Replace with your video file path
    audio_text = audio_to_text(video_file)
    frame_text = video_to_text(video_file)
    create_document(audio_text, frame_text)
